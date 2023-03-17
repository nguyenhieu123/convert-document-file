from docx import Document
from docx.shared import Inches, Pt, RGBColor

def make_client_invoice(name, email, product, unit, price):
    documnet = Document()
    documnet.add_heading('Invoice', 0)
    p1 = documnet.add_paragraph('Dear ')
    p1.add_run(name).bold = True
    p1.add_run(',')

    p2 = documnet.add_paragraph('Please find attached invoice for your recent purchase of ')
    p2.add_run(str(unit)).bold = True
    p2.add_run(' unit of ')
    p2.add_run(product).bold = True
    p2.add_run('.')

    [documnet.add_paragraph('') for _ in range(2)]

    table = documnet.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product Name'
    hdr_cells[1].text = 'Units'
    hdr_cells[2].text = 'Units Price'
    hdr_cells[3].text = 'Total Price'
    # for i in range(4):
    #     hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    row_cells = table.add_row().cells
    row_cells[0].text = product
    row_cells[1].text = f'{unit:,.2f}'
    row_cells[2].text = f'{price:,.2f}'
    row_cells[3].text = f'{unit*price:,.2f}'

    [documnet.add_paragraph('') for _ in range(10)]
    documnet.add_paragraph('We appreciate your business and please come again!')
    documnet.add_paragraph('Jay')
    documnet.save(f'{name}1.docx')

make_client_invoice('Tom', 'abc@gmail.com', 'mobile', 10, 1000)

